import pandas as pd
from docx import Document
import os
import win32api
import win32print
import time 
from datetime import datetime
import locale
import comtypes.client 

# Configurações

ARQUIVO_EXCEL = "dados.xlsx"         # Planilha com os dados de entrada
NOME_ABA = "revisoes"                # Nome da aba no Excel
TEMPLATE_WORD = "template.docx"      # Template base
PASTA_SAIDA = "relatorios_gerados"   # Pasta onde os .docx serão salvos
CIDADE = "São Vicente"


# Configuração data e local
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
except:
    locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')  # Windows

data_por_extenso = datetime.now().strftime("%d de %B de %Y")
linha_data_cidade = f"{CIDADE}, {data_por_extenso}"


# 1. Criar a pasta de saída se não existir
if not os.path.exists(PASTA_SAIDA):
    os.makedirs(PASTA_SAIDA)

# 2. Ler o arquivo Excel
try:
    # Ler o arquivo Excel
    df = pd.read_excel(ARQUIVO_EXCEL, sheet_name=NOME_ABA, header=0)
    
    # 3. Formatar as colunas de data para o formato datetime
    df['DATA REALIZADA'] = pd.to_datetime(df['DATA REALIZADA'], errors='coerce')
    df['PROXIMA DATA'] = pd.to_datetime(df['PROXIMA DATA'], errors='coerce')
    
except FileNotFoundError:
    print(f"ERRO: Arquivo '{ARQUIVO_EXCEL}' não encontrado.")
    exit()
except Exception as e:
    print(f"ERRO ao ler ou processar dados: {e}")
    exit()


# 4. Agrupar os dados por  ID (PREFIXO)
df_agrupado = df.groupby('PREFIXO')

print(f"Encontrados {len(df_agrupado)} prefixos únicos. Gerando documentos...")

# 5. Loop para cada prefixo
for prefixo, group_df in df_agrupado:
    
    print(f"Processando Prefixo: {prefixo}...")
    
    try:
        doc = Document(TEMPLATE_WORD)
    except Exception as e:
        print(f"ERRO: Não foi possível abrir o template '{TEMPLATE_WORD}'. {e}")
        continue

    if not doc.tables:
        print(f"ERRO: O template '{TEMPLATE_WORD}' não contém tabelas.")
        continue
        
    tabela = doc.tables[0]


   # Tenta primeiro no cabeçalho do documento
    inserido = False
    for section in doc.sections:
        header = section.header
        if header.is_linked_to_previous == False or len(doc.sections) == 1:
            # Adiciona como primeiro parágrafo do cabeçalho
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = linha_data_cidade
            p.style = 'Header' 
            p.paragraph_format.alignment = 1  # centralizado
            inserido = True
            break

    # Se não conseguiu no cabeçalho, coloca como primeiro parágrafo do corpo
    if not inserido:
        p = doc.paragraphs[0]
        p.text = linha_data_cidade
        p.style = 'Hearder'
        p.paragraph_format.alignment = 1  # centralizado

    # 6. Loop nas linhas de dados
    for _, linha in group_df.iterrows():
        
        # Formatação das Datas (DD/MM/AAAA)
        data_higi_formatada = linha['DATA REALIZADA'].strftime('%d/%m/%Y') if pd.notna(linha['DATA REALIZADA']) else ''
        data_prox_formatada = linha['PROXIMA DATA'].strftime('%d/%m/%Y') if pd.notna(linha['PROXIMA DATA']) else '' 

        # Adiciona uma nova linha à tabela no Word
        nova_linha_celulas = tabela.add_row().cells
        
        # Preenche as células com as datas formatadas
        nova_linha_celulas[0].text = str(linha['PREFIXO'])
        nova_linha_celulas[1].text = str(linha['EMPRESA'])
        nova_linha_celulas[2].text = data_higi_formatada # Usa a string formatada
        nova_linha_celulas[3].text = data_prox_formatada # Usa a string formatada

    # 7. Salvar e Imprimir
    nome_arquivo = os.path.join(PASTA_SAIDA, f"Relatorio_{prefixo}.pdf")
    
    # Converte para PDF usando Word
    try:
        temp_docx = os.path.join(PASTA_SAIDA, f"temp_{prefixo}.docx")
        doc.save(temp_docx)

        
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc_com = word.Documents.Open(os.path.abspath(temp_docx))
        doc_com.SaveAs(os.path.abspath(nome_arquivo), FileFormat=17) 
        doc_com.Close()
        word.Quit()
        os.remove(temp_docx)
        print(f"   PDF gerado → Relatorio_{prefixo}.pdf")
    except Exception as e:
        print(f"   Erro ao gerar PDF (Word não encontrado?): {e}")
        # fallback
        doc.save(os.path.join(PASTA_SAIDA, f"Relatorio_{prefixo}.docx"))

    """# Bloco de impressão
    try:
        print(f"Enviando '{nome_arquivo_saida}' para a impressora...")
        caminho_completo = os.path.abspath(nome_arquivo_saida)
        win32api.ShellExecute(
            0,
            "print",
            caminho_completo,
            f'/d:"{win32print.GetDefaultPrinter()}"',
            ".",
            0
        )
        print("Enviado.")
        time.sleep(2) 
    except Exception as e:
        print(f"ERRO ao tentar imprimir: {e}")"""

print(f"\nConcluído! {len(df_agrupado)} documentos foram gerados na pasta '{PASTA_SAIDA}'.")